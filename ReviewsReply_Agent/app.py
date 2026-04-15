import os
import json
import re
import datetime
import urllib.request
from flask import Flask, request, jsonify, render_template
from docx import Document

app = Flask(__name__)

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL = "mistral"

# -----------------------------
# HELPERS
# -----------------------------

def call_ollama(prompt):
    data = json.dumps({
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.3,
            "num_predict": 180
        }
    }).encode("utf-8")

    req = urllib.request.Request(
        OLLAMA_URL,
        data=data,
        headers={'Content-Type': 'application/json'}
    )

    try:
        with urllib.request.urlopen(req, timeout=20) as res:
            return json.loads(res.read().decode())["response"]
    except:
        return "Error connecting to Ollama."


def extract_first_name(full_name):
    return full_name.split()[0] if full_name else "Guest"


def capitalize_name(name):
    return " ".join([w.capitalize() for w in name.split()])


def extract_narrative(review_text, reviewer_name):
    lines = review_text.split('\n')

    if lines and lines[0].strip().lower() == reviewer_name.lower():
        lines = lines[1:]

    clean_lines = []
    for line in lines:
        l = line.lower().strip()

        if re.search(r'\d+\s+review', l):
            continue
        if "photo" in l:
            continue
        if any(x in l for x in ["dine in", "price", "₹", "$", "food:", "service:", "atmosphere:", "wait time"]):
            continue
        if "star" in l or "⭐" in l:
            continue
        if not l or len(l) < 2:
            continue

        clean_lines.append(line.strip())

    return " ".join(clean_lines).strip()


# -----------------------------
# DOCX HANDLING
# -----------------------------

def remove_existing_review(doc, reviewer_name):
    paragraphs = doc.paragraphs
    indices_to_remove = []

    for i, p in enumerate(paragraphs):
        if p.text.strip() == reviewer_name:
            indices_to_remove.extend([i, i+1, i+2])

    for i in sorted(indices_to_remove, reverse=True):
        try:
            p = paragraphs[i]._element
            p.getparent().remove(p)
        except:
            pass


def save_review(restaurant, reviewer, review_text):
    now = datetime.datetime.now()
    month, year = now.strftime("%B"), now.strftime("%Y")

    prefix = "Aroma" if "Aroma" in restaurant else "Urban"
    filename = f"{prefix}_{month}_{year}.docx"

    reviewer_cap = capitalize_name(reviewer)
    narrative = extract_narrative(review_text, reviewer_cap)

    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
        doc.add_heading(f"{prefix} Reviews - {month} {year}", 0)

    remove_existing_review(doc, reviewer_cap)

    if len(doc.paragraphs) > 1:
        ref = doc.paragraphs[1]
        ref.insert_paragraph_before("-" * 50)
        ref.insert_paragraph_before(narrative)
        ref.insert_paragraph_before(reviewer_cap)
    else:
        doc.add_paragraph(reviewer_cap)
        doc.add_paragraph(narrative)
        doc.add_paragraph("-" * 50)

    doc.save(filename)


# -----------------------------
# BRAND CONTROL
# -----------------------------

def enforce_brand_rules(reply, restaurant, signoff):

    # Remove PS lines
    reply = re.sub(r"P\.S\.:.*", "", reply, flags=re.IGNORECASE)

    if "Aroma" in restaurant:
        reply = re.sub(r"Aroma Indian Cuisine(?! \| Bar)", "Aroma Indian Cuisine | Bar", reply)
        reply = re.sub(r"Urban Indian", "", reply)
    else:
        reply = re.sub(r"Aroma Indian Cuisine.*", "", reply)

    # Remove any incorrect signoffs
    reply = re.sub(r"Warm regards,.*", "", reply, flags=re.DOTALL)

    reply = reply.strip() + f"\n\n{signoff}"

    return reply.strip()


# -----------------------------
# ROUTES
# -----------------------------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/vault')
def vault_page():
    return render_template('vault.html')


@app.route('/readymade')
def readymade_page():
    return render_template('readymade.html')


# -----------------------------
# API
# -----------------------------

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json
    review = data.get("review_text", "")
    restaurant = data.get("restaurant", "Urban Indian")

    full_name = review.split('\n')[0] if review else "Guest"
    full_name = capitalize_name(full_name)

    first_name = extract_first_name(full_name)

    if "Aroma" in restaurant:
        signoff = "Warm regards,\nAroma Indian Cuisine | Bar Team."
        brand_instruction = "Use 'Aroma Indian Cuisine | Bar' whenever mentioning the restaurant."
    else:
        signoff = "Warm regards,\nTeam Urban Indian."
        brand_instruction = "Use 'Urban Indian' whenever mentioning the restaurant."

    prompt = f"""
Write a professional restaurant reply.

STRICT RULES:
- USE ONLY first name: {first_name}
- START THE REPLY IMMEDIATELY with this sentence: Thank you for taking the time to share your positive experience at {restaurant}, {first_name}!
- DO NOT use any salutation (NO 'Dear', 'Hi', 'Hello', etc.)
- NO headings
- NO meta commentary
- NO P.S.
- NO extra names
- ONLY 2 paragraphs
- Keep it natural and human
- {brand_instruction}

End EXACTLY with:
{signoff}

Review:
{review}
"""

    reply = call_ollama(prompt)
    reply = enforce_brand_rules(reply, restaurant, signoff)

    save_review(restaurant, full_name, review)

    return jsonify({"reply": reply})


@app.route('/api/readymade', methods=['POST'])
def readymade_api():
    data = request.json
    restaurant = data.get("restaurant", "Urban Indian")
    category = data.get("type", "5")

    if "Aroma" in restaurant:
        signoff = "Warm regards,\nAroma Indian Cuisine | Bar Team."
    else:
        signoff = "Warm regards,\nTeam Urban Indian."

    prompt = f"""
Generate a {category}-star restaurant reply.

Rules:
- Use placeholder [Customer Name]
- START THE REPLY IMMEDIATELY with this sentence: Thank you for taking the time to share your positive experience at {restaurant}, [Customer Name]!
- DO NOT use any salutation (NO 'Dear', 'Hi', 'Hello', etc.)
- 2 paragraphs only
- No headings
- No extra lines
- End EXACTLY with:

{signoff}
"""

    reply = call_ollama(prompt)
    reply = enforce_brand_rules(reply, restaurant, signoff)

    return jsonify({"reply": reply})


# -----------------------------
# UPDATED VAULT LOGIC (YOUR CHANGE)
# -----------------------------

@app.route('/api/vault-files')
def vault_files():
    import datetime
    import os

    now = datetime.datetime.now()
    current_month = now.strftime("%B")
    current_year = now.strftime("%Y")

    active = []
    history = []

    for f in os.listdir('.'):
        if f.endswith('.docx') and not f.startswith('~$'):
            if current_month in f and current_year in f:
                active.append(f)
            else:
                history.append(f)

    return jsonify({
        "active": sorted(active),
        "history": sorted(history, reverse=True)
    })


@app.route('/api/file/<filename>')
def get_file_content(filename):
    if not os.path.exists(filename):
        return jsonify([])

    doc = Document(filename)
    return jsonify([p.text for p in doc.paragraphs if p.text.strip()])


# -----------------------------
# RUN
# -----------------------------

if __name__ == "__main__":
    app.run(debug=True)