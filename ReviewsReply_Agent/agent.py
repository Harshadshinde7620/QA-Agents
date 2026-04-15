import os
import json
import urllib.request
import urllib.error
from docx import Document
from dotenv import load_dotenv

OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_MODEL = "mistral"

def initialize_api():
    """Checks that the Ollama server is reachable using urllib."""
    load_dotenv()
    try:
        url = f"{OLLAMA_BASE_URL}/api/tags"
        with urllib.request.urlopen(url, timeout=5) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            models = [m["name"] for m in data.get("models", [])]
            mistral_found = any(OLLAMA_MODEL in m for m in models)
            if not mistral_found:
                print(f"WARNING: '{OLLAMA_MODEL}' model not found on Ollama server.")
                print(f"Available models: {models}")
                print(f"Run: ollama pull {OLLAMA_MODEL}")
                return False
            print(f"✅ Ollama server connected. Using model: {OLLAMA_MODEL}")
            return True
    except Exception as e:
        print("ERROR: Cannot connect to Ollama. Make sure Ollama is running.")
        print("Start it with: ollama serve")
        print(f"Technical error: {e}")
        return False

def generate_reply_and_extract(username, review_text, restaurant_name, restaurant_full_name):
    """Uses Ollama/Mistral to generate a reply and extract key phrases using urllib."""

    prompt = f"""You are a professional reputation manager for a restaurant named '{restaurant_name}' (Full name: '{restaurant_full_name}').

A customer named '{username}' just left the following review:
"{review_text}"

Please provide me with a JSON response containing two things:
1. "reply": A professional, polite reply to the review. The reply MUST include the name '{username}'. At the end of the reply, sign off with regards along with the full name of the restaurant ('{restaurant_full_name}').
2. "extracted_phrases": A list of the actual wordings/key phrases inputted by the reviewer expressing their opinion on the food, service, atmosphere, etc. (e.g., "fantastic food", "excellent service"). Keep the exact wordings used by the reviewer.

Output strictly as JSON in the following format, with no markdown formatting around it:
{{
    "reply": "Dear [username], ... Warm regards, [restaurant_full_name]",
    "extracted_phrases": ["phrase 1", "phrase 2"]
}}"""

    url = f"{OLLAMA_BASE_URL}/api/generate"
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "format": "json"
    }
    encoded_data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(url, data=encoded_data, headers={'Content-Type': 'application/json'})

    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            data = json.loads(result["response"])
            return data.get("reply"), data.get("extracted_phrases", [])
    except Exception as e:
        print(f"Error during Ollama API call: {e}")
        return None, []

def save_to_word_file(username, extracted_phrases, filename="Review_Keywords.docx"):
    """Saves the extracted key phrases to a Word document."""
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
        doc.add_heading('Customer Review Key Phrases', 0)

    doc.add_paragraph(f"Reviewer: {username}")

    if extracted_phrases:
        p = doc.add_paragraph("Key phrases: ")
        p.add_run(", ".join(extracted_phrases)).italic = True
    else:
        doc.add_paragraph("Key phrases: None identified")

    doc.add_paragraph("-" * 40)
    doc.save(filename)
    print(f"Saved extracted phrases to '{filename}'")

def main():
    print("=" * 50)
    print("Welcome to the Restaurant Review Reply Agent")
    print("         Powered by Ollama + Mistral")
    print("=" * 50)

    if not initialize_api():
        return

    # Basic Configuration
    restaurant_name = input("Enter your restaurant's short name (e.g., The Grand Eatery): ").strip()
    restaurant_full_name = input("Enter your restaurant's full name (e.g., The Grand Eatery & Fine Dining): ").strip()

    if not restaurant_name:
        restaurant_name = "Our Restaurant"
    if not restaurant_full_name:
        restaurant_full_name = "Our Restaurant"

    while True:
        print("\n" + "-" * 30)
        username = input("Enter the reviewer's name (or type 'exit' to quit): ").strip()
        if username.lower() == 'exit':
            break

        review_text = input(f"Enter the review provided by {username}: ").strip()
        if not review_text:
            print("Review text cannot be empty. Try again.")
            continue

        print("\nGenerating reply and extracting phrases with Ollama (Mistral)...")
        reply, extracted_phrases = generate_reply_and_extract(username, review_text, restaurant_name, restaurant_full_name)

        if reply:
            print("\n" + "=" * 40)
            print("GENERATED REVIEW REPLY\n")
            print(reply)
            print("=" * 40 + "\n")

            print(f"Extracted Phrases: {extracted_phrases}")
            save_to_word_file(username, extracted_phrases)
        else:
            print("Failed to generate a reply. Please check Ollama status.")

if __name__ == "__main__":
    main()
